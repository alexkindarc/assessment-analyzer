"""
Assessment Report Analyzer v3
Comprehensive tool for analyzing academic and administrative assessment reports
with metadata tracking, stagnation detection, and multi-report type support.
"""

import streamlit as st
import anthropic
import pandas as pd
from datetime import datetime
import json
import os
import re
from io import BytesIO
import hashlib

# Optional imports for file handling
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# Microsoft Graph API support for Excel Online
try:
    import msal
    import requests
    EXCEL_ONLINE_SUPPORT = True
except ImportError:
    EXCEL_ONLINE_SUPPORT = False

# Page configuration
st.set_page_config(
    page_title="Assessment Report Analyzer | UTA",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================================================
# UTA BRAND STYLING
# ============================================================================

UTA_CSS = """
<style>
    /* Clean Modern Design - UTA Colors */
    :root {
        --uta-blue: #0064b1;
        --uta-dark-blue: #003865;
        --uta-orange: #F58025;
        --sidebar-bg: #f0f4f8;
        --content-bg: #ffffff;
        --border-color: #e1e5eb;
        --text-primary: #1a2b3c;
        --text-secondary: #5a6777;
        --text-muted: #8896a6;
    }
    
    /* Hide Streamlit branding but KEEP navigation controls */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    
    /* IMPORTANT: Make sure sidebar toggle is always visible */
    button[kind="header"] {
        visibility: visible !important;
        display: block !important;
    }
    
    /* Remove top padding */
    .block-container {
        padding-top: 2rem;
        padding-bottom: 1rem;
        max-width: 100%;
    }
    
    /* Sidebar - Light clean design with accent */
    section[data-testid="stSidebar"] {
        background-color: var(--sidebar-bg);
        border-right: 1px solid var(--border-color);
    }
    
    section[data-testid="stSidebar"] > div {
        padding-top: 1rem;
    }
    
    /* Sidebar text colors - dark on light */
    section[data-testid="stSidebar"] .stMarkdown,
    section[data-testid="stSidebar"] p,
    section[data-testid="stSidebar"] span,
    section[data-testid="stSidebar"] label {
        color: var(--text-primary) !important;
    }
    
    section[data-testid="stSidebar"] h1,
    section[data-testid="stSidebar"] h2,
    section[data-testid="stSidebar"] h3 {
        color: var(--uta-dark-blue) !important;
    }
    
    /* Sidebar inputs */
    section[data-testid="stSidebar"] .stTextInput input,
    section[data-testid="stSidebar"] .stSelectbox > div > div {
        background-color: white !important;
        border: 1px solid var(--border-color) !important;
        color: var(--text-primary) !important;
    }
    
    /* Sidebar buttons - ensure readable text */
    section[data-testid="stSidebar"] .stButton > button {
        background-color: white !important;
        border: 1px solid #e1e5eb !important;
        color: #1a2b3c !important;
        border-radius: 8px !important;
        font-weight: 500 !important;
        text-align: left !important;
        padding: 0.75rem 1rem !important;
    }
    
    section[data-testid="stSidebar"] .stButton > button:hover {
        background-color: #f0f4f8 !important;
        border-color: #0064b1 !important;
        color: #0064b1 !important;
    }
    
    /* Primary buttons in sidebar (active nav) */
    section[data-testid="stSidebar"] .stButton > button[kind="primary"],
    section[data-testid="stSidebar"] .stButton > button[data-testid="baseButton-primary"] {
        background-color: #0064b1 !important;
        border-color: #0064b1 !important;
        color: white !important;
    }
    
    /* Disabled primary buttons (current page indicator) */
    section[data-testid="stSidebar"] .stButton > button[kind="primary"]:disabled,
    section[data-testid="stSidebar"] .stButton > button[data-testid="baseButton-primary"]:disabled {
        background-color: #0064b1 !important;
        border-color: #0064b1 !important;
        color: white !important;
        opacity: 1 !important;
        cursor: default !important;
    }
    
    /* Fix text inside all sidebar buttons */
    section[data-testid="stSidebar"] .stButton > button p,
    section[data-testid="stSidebar"] .stButton > button span,
    section[data-testid="stSidebar"] .stButton > button div {
        color: inherit !important;
    }
    
    /* Sidebar dividers */
    section[data-testid="stSidebar"] hr {
        border-color: var(--border-color) !important;
        margin: 1rem 0 !important;
    }
    
    /* Sidebar expander */
    section[data-testid="stSidebar"] .streamlit-expanderHeader {
        background-color: white !important;
        border: 1px solid var(--border-color) !important;
        border-radius: 8px !important;
        color: var(--text-primary) !important;
    }
    
    section[data-testid="stSidebar"] .streamlit-expanderContent {
        background-color: white !important;
        border: 1px solid var(--border-color) !important;
        border-top: none !important;
        border-radius: 0 0 8px 8px !important;
    }
    
    /* Status badges in sidebar */
    section[data-testid="stSidebar"] .stAlert {
        background-color: white !important;
        border: 1px solid var(--border-color) !important;
        border-radius: 8px !important;
        padding: 0.5rem 0.75rem !important;
    }
    
    /* Main content area */
    .main .block-container {
        background-color: var(--content-bg);
        color: var(--text-primary);
    }
    
    /* Page title styling with orange accent */
    .page-title {
        color: var(--uta-dark-blue);
        font-size: 1.5rem;
        font-weight: 600;
        margin-bottom: 0.25rem;
        padding-bottom: 0.75rem;
        border-bottom: 3px solid var(--uta-orange);
        display: inline-block;
    }
    
    /* Section headers with accent */
    .section-header {
        color: var(--uta-dark-blue);
        font-size: 0.8rem;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        margin-bottom: 0.75rem;
        margin-top: 1.5rem;
        padding-left: 0.5rem;
        border-left: 3px solid var(--uta-orange);
    }
    
    /* Content cards */
    .content-card {
        background-color: white;
        border: 1px solid var(--border-color);
        border-radius: 12px;
        padding: 1.5rem;
        margin-bottom: 1rem;
    }
    
    /* Form field styling */
    .stTextInput > label,
    .stSelectbox > label,
    .stTextArea > label {
        color: var(--text-primary) !important;
        font-weight: 500 !important;
        font-size: 0.875rem !important;
    }
    
    .stTextInput input,
    .stTextArea textarea {
        border: 1px solid var(--border-color) !important;
        border-radius: 8px !important;
        padding: 0.625rem 0.875rem !important;
        color: var(--text-primary) !important;
    }
    
    /* FIX SELECTBOX TEXT VISIBILITY */
    .stSelectbox > div > div {
        border: 1px solid var(--border-color) !important;
        border-radius: 8px !important;
        background-color: white !important;
    }
    
    .stSelectbox > div > div > div {
        color: var(--text-primary) !important;
    }
    
    .stSelectbox [data-baseweb="select"] > div {
        background-color: white !important;
        color: var(--text-primary) !important;
    }
    
    .stSelectbox [data-baseweb="select"] span {
        color: var(--text-primary) !important;
    }
    
    /* Dropdown menu items */
    [data-baseweb="popover"] {
        background-color: white !important;
    }
    
    [data-baseweb="popover"] li {
        color: var(--text-primary) !important;
    }
    
    [data-baseweb="popover"] li:hover {
        background-color: var(--sidebar-bg) !important;
    }
    
    .stTextInput input:focus,
    .stSelectbox > div > div:focus,
    .stTextArea textarea:focus {
        border-color: var(--uta-blue) !important;
        box-shadow: 0 0 0 3px rgba(0, 100, 177, 0.1) !important;
    }
    
    /* Primary button with orange accent on hover */
    .stButton > button[kind="primary"] {
        background-color: var(--uta-blue) !important;
        border: none !important;
        border-radius: 8px !important;
        color: white !important;
        font-weight: 500 !important;
        padding: 0.625rem 1.25rem !important;
        transition: all 0.2s ease !important;
    }
    
    .stButton > button[kind="primary"]:hover {
        background-color: var(--uta-dark-blue) !important;
        box-shadow: 0 4px 12px rgba(0, 100, 177, 0.3) !important;
    }
    
    /* Secondary button */
    .stButton > button[kind="secondary"],
    .stDownloadButton > button {
        background-color: white !important;
        border: 1px solid var(--border-color) !important;
        border-radius: 8px !important;
        color: var(--text-primary) !important;
        font-weight: 500 !important;
    }
    
    .stButton > button[kind="secondary"]:hover,
    .stDownloadButton > button:hover {
        background-color: var(--sidebar-bg) !important;
        border-color: var(--uta-blue) !important;
    }
    
    /* File uploader with accent */
    [data-testid="stFileUploader"] {
        background-color: white;
        border: 2px dashed var(--border-color);
        border-radius: 12px;
        padding: 2rem;
        transition: all 0.2s ease;
    }
    
    [data-testid="stFileUploader"]:hover {
        border-color: var(--uta-orange);
        background-color: rgba(245, 128, 37, 0.02);
    }
    
    /* Success/Info/Warning alerts with accents */
    .stSuccess {
        background-color: #f0fdf4 !important;
        border: 1px solid #86efac !important;
        border-left: 4px solid #22c55e !important;
        border-radius: 8px !important;
        color: #166534 !important;
    }
    
    .stInfo {
        background-color: #eff6ff !important;
        border: 1px solid #93c5fd !important;
        border-left: 4px solid var(--uta-blue) !important;
        border-radius: 8px !important;
        color: #1e40af !important;
    }
    
    .stWarning {
        background-color: #fffbeb !important;
        border: 1px solid #fcd34d !important;
        border-left: 4px solid var(--uta-orange) !important;
        border-radius: 8px !important;
        color: #92400e !important;
    }
    
    .stError {
        background-color: #fef2f2 !important;
        border: 1px solid #fca5a5 !important;
        border-left: 4px solid #ef4444 !important;
        border-radius: 8px !important;
        color: #991b1b !important;
    }
    
    /* Expander in main content */
    .main .streamlit-expanderHeader {
        background-color: var(--sidebar-bg) !important;
        border: 1px solid var(--border-color) !important;
        border-radius: 8px !important;
        color: var(--text-primary) !important;
        font-weight: 500 !important;
    }
    
    .main .streamlit-expanderContent {
        border: 1px solid var(--border-color) !important;
        border-top: none !important;
        border-radius: 0 0 8px 8px !important;
        background-color: white !important;
    }
    
    /* Tabs styling with orange accent */
    .stTabs [data-baseweb="tab-list"] {
        gap: 0;
        background-color: var(--sidebar-bg);
        border-radius: 8px;
        padding: 4px;
    }
    
    .stTabs [data-baseweb="tab"] {
        border-radius: 6px;
        padding: 8px 16px;
        color: var(--text-secondary) !important;
        font-weight: 500;
        background-color: transparent;
    }
    
    .stTabs [aria-selected="true"] {
        background-color: white !important;
        color: var(--uta-blue) !important;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        border-bottom: 2px solid var(--uta-orange) !important;
    }
    
    /* Divider */
    hr {
        border: none;
        border-top: 1px solid var(--border-color);
        margin: 1.5rem 0;
    }
    
    /* Footer with accent */
    .app-footer {
        text-align: center;
        padding: 1.5rem;
        color: var(--text-muted);
        font-size: 0.8rem;
        border-top: 1px solid var(--border-color);
        margin-top: 2rem;
        background: linear-gradient(to right, var(--uta-blue), var(--uta-dark-blue));
        background-size: 100% 3px;
        background-repeat: no-repeat;
        background-position: top;
    }
    
    /* Analysis results card with accent */
    .results-card {
        background-color: white;
        border: 1px solid var(--border-color);
        border-left: 4px solid var(--uta-blue);
        border-radius: 8px;
        padding: 1.25rem;
        margin-bottom: 1rem;
    }
    
    /* Metadata field groups */
    .field-group {
        background-color: var(--sidebar-bg);
        border-radius: 8px;
        padding: 1rem;
        margin-bottom: 1rem;
        border-left: 3px solid var(--uta-orange);
    }
    
    .field-group-title {
        color: var(--text-muted);
        font-size: 0.7rem;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        margin-bottom: 0.75rem;
    }
    
    /* Outcome card with accent */
    .outcome-card {
        background-color: white;
        border: 1px solid var(--border-color);
        border-left: 3px solid var(--uta-blue);
        border-radius: 8px;
        padding: 1rem;
        margin-bottom: 0.75rem;
    }
    
    /* Spinner */
    .stSpinner > div {
        border-color: var(--uta-blue) !important;
    }
    
    /* Caption text */
    .stCaption {
        color: var(--text-muted) !important;
    }
    
    /* Logo area with accent bar */
    .logo-area {
        padding: 1rem 0.5rem 1.5rem 0.5rem;
        border-bottom: 1px solid var(--border-color);
        margin-bottom: 1rem;
        position: relative;
    }
    
    .logo-area::after {
        content: '';
        position: absolute;
        bottom: -1px;
        left: 0;
        width: 60px;
        height: 3px;
        background-color: var(--uta-orange);
    }
    
    .logo-text {
        color: var(--uta-dark-blue);
        font-size: 1.25rem;
        font-weight: 700;
        line-height: 1.3;
    }
    
    /* Connection status pill */
    .status-pill {
        display: inline-flex;
        align-items: center;
        padding: 0.25rem 0.75rem;
        border-radius: 999px;
        font-size: 0.75rem;
        font-weight: 500;
    }
    
    .status-connected {
        background-color: #dcfce7;
        color: #166534;
    }
    
    .status-pending {
        background-color: #fef3c7;
        color: #92400e;
    }
    
    .status-disconnected {
        background-color: #f3f4f6;
        color: #6b7280;
    }
    
    /* Accent line for visual interest */
    .accent-line {
        width: 60px;
        height: 3px;
        background-color: var(--uta-orange);
        margin: 1rem 0;
    }
</style>
"""

def inject_uta_branding():
    """Inject UTA branding CSS into the page."""
    st.markdown(UTA_CSS, unsafe_allow_html=True)

def render_uta_header(title: str, subtitle: str = None):
    """Render clean page header."""
    st.markdown(f'<h1 class="page-title">{title}</h1>', unsafe_allow_html=True)

def render_uta_footer():
    """Render clean footer."""
    footer_html = """
    <div class="app-footer">
        <strong>Assessment Report Analyzer</strong> · Office of Institutional Effectiveness and Reporting · The University of Texas at Arlington
    </div>
    """
    st.markdown(footer_html, unsafe_allow_html=True)

# ============================================================================
# CONSTANTS AND CONFIGURATION
# ============================================================================

STRATEGIC_THEMES = [
    "People and Culture",
    "Student Success", 
    "Alumni and Community Engagement",
    "Research and Innovation",
    "Finance and Infrastructure"
]

CORE_OBJECTIVES = [
    "Critical Thinking",
    "Communication",
    "Empirical & Quantitative Skills",
    "Teamwork",
    "Personal Responsibility",
    "Social Responsibility"
]

ACHIEVEMENT_LEVELS = [
    "Fully Achieved",      # Met ALL criteria for success
    "Partially Achieved",  # Multiple criteria exist; some met, some not
    "Not Achieved",        # Did not meet criteria for success
    "Inconclusive"         # Insufficient data or unclear results (must be explained)
]

REPORT_TYPES = [
    "Results Report",
    "Improvement Report", 
    "Next Cycle Plan"
]

# ============================================================================
# DEFAULT PROMPTS - Editable by Admin
# ============================================================================

DEFAULT_RUBRIC_GUIDANCE = """## Assessment Report Quality Guidance

This rubric guides analysis and feedback. It does NOT assign scores or grades.
Reference these criteria when providing feedback, explaining WHY each element matters.

### Student Learning Outcomes / Outcomes
- Should be specific and measurable
- Should use appropriate action verbs
- Should align with program/institutional goals
- For academic programs: labeled "Student Learning Outcomes"
- For administrative units: labeled "Outcomes" (never "learning outcomes")

### Assessment Methods
- Should use direct measures where possible
- Should align clearly with each outcome
- Should describe instruments/rubrics adequately

### Data Collection & Reporting
- Sample size should be adequate for conclusions
- Results MUST match the format of criteria for success
- If benchmark says "75% of students score 90%+", results must report in that exact format
- Vague results like "most students passed" are insufficient

### Benchmarks/Criteria for Success
- Should be clearly stated before results
- Should be appropriately ambitious yet achievable

### Use of Results (Closing the Loop)
- Improvements should be SPECIFIC, not generic
- Should directly address identified weaknesses
- Should identify responsible parties and timelines
- Different outcomes should have differentiated improvements

### Strategic Alignment
- Should map to Strategic Plan themes where applicable
- Should connect to program competencies or core functions
"""

DEFAULT_TONE_INSTRUCTIONS = """## Communication Tone

Write as a supportive colleague, not an auditor or critic.

REQUIRED APPROACH:
- Begin with genuine acknowledgment of what's working well
- Frame gaps as "opportunities to strengthen" not "failures" or "deficiencies"
- Use phrases like "consider," "you might," "one option would be"
- Avoid "you must," "you failed to," "this is wrong"
- Assume good faith - units are trying to improve
- Offer specific, actionable suggestions alongside any concerns
- End on a constructive, encouraging note

EXAMPLE TRANSFORMATION:
Instead of: "The improvement actions are generic and fail to address specific findings."
Write: "The improvement actions provide a solid starting point. To strengthen this section, consider adding specifics about which particular topics students found challenging and how instruction will address those gaps. This helps demonstrate the direct connection between your findings and your response."

Remember: The goal is to help units improve their assessment practice, not to penalize them.
"""

DEFAULT_RESULTS_ANALYSIS_PROMPT = """You are an expert in higher education assessment, supporting colleagues in improving their assessment practice. You have deep knowledge of SACSCOC, AACSB, and other accreditation standards.

{rubric_guidance}

{tone_instructions}

## Your Task: Analyze this Results Report

Review the assessment report and run ALL of the following checks internally. Do NOT output each check separately. Instead, synthesize your findings into a concise, actionable response.

### INTERNAL CHECKS TO RUN (do not output these sections):

1. **Field Completeness** - All required fields filled (Outcome Rationale is optional)

2. **Competencies/Functions** - Must be coherent statements (not single words), and each outcome's Related Competency/Function must match one listed in General Information

3. **Bloom's Taxonomy** - Action verbs appropriate for program level:
   - UG lower: Levels 1-3 OK
   - UG upper/capstone: Levels 3-5 expected
   - Graduate: Levels 4-6 expected (1-3 too low)
   - Doctoral: Levels 5-6 expected
   - Flag vague verbs: "understand," "know," "learn," "appreciate"

4. **Results-Criteria Alignment** - Results must align with criteria for success IN SUBSTANCE (not verbatim). If criteria says "75% score 80%+", results should report a percentage against that threshold. Vague results like "most students did well" are insufficient.

5. **Quantitative Data** - Methodology must produce at least one quantitative result

6. **Sample Size** - When percentages reported, sample size (n=) must be provided

7. **Achievement Level Logic**:
   - Fully Achieved: Results meet ALL criteria
   - Partially Achieved: ONLY valid with multiple criteria where some met, some not
   - Not Achieved: Results don't meet criteria
   - Inconclusive: Must be explained in report

8. **Action Steps** - Must be specific (who, what, when), not vague or copy-pasted across outcomes

9. **Proposed Improvements** - For outcomes not achieved, must have sufficient detail (what will change, who responsible, timeline), not just "we will try harder"

{stagnation_context}

{custom_rubric}

## YOUR OUTPUT FORMAT:

Provide a concise response with TWO sections only:

### ✓ STRENGTHS
2-3 bullet points of what the report does well. Be genuine and specific.

### ✎ REVISIONS REQUESTED

CRITICAL REQUIREMENT: Every single revision bullet MUST begin with the outcome number in bold. No exceptions.

Format each bullet EXACTLY like this:
- **Outcome 1:** [specific issue and how to fix it]
- **Outcome 2:** [specific issue and how to fix it]  
- **Outcomes 1, 2, 3:** [if the same issue applies to multiple outcomes]
- **All outcomes:** [only if truly applies to every outcome]

DO NOT write generic bullets without outcome numbers. If a revision applies to a specific outcome, identify it by number.

Example of CORRECT format:
- **Outcome 1:** Results report an average score (78%) but criteria states "75% of students will score 80%+". Please report what percentage of students met the 80% threshold.
- **Outcome 3:** Sample size missing. Add n= to show how many students were assessed.
- **Outcomes 2 & 4:** Action steps are identical. Differentiate the improvement plans based on each outcome's specific findings.

Example of WRONG format (do not do this):
- Consider revising outcomes to use more specific action verbs. ← WRONG: doesn't specify which outcome
- Add sample sizes where missing. ← WRONG: doesn't specify which outcome

If no revisions needed, say "No revisions needed - report meets all criteria."

Keep your total response under 500 words. Be direct and actionable.

## Report to Analyze:

{report_text}
"""

DEFAULT_IMPROVEMENT_ANALYSIS_PROMPT = """You are an expert in higher education assessment, supporting colleagues in documenting their improvement efforts. 

{rubric_guidance}

{tone_instructions}

## Your Task: Analyze this Improvement Report

This report documents actions the unit has taken based on previous assessment findings. Run checks internally and provide a concise response.

### INTERNAL CHECKS TO RUN (do not output these sections):

1. **Clarity of Actions** - Are improvement actions clearly described with enough detail to demonstrate genuine effort?

2. **Connection to Previous Findings** - Does the report reference what was originally found and how actions address it? (This is good practice but not strictly required - note gently if missing)

3. **Specificity** - Do actions describe what specifically was done, not just vague statements like "we improved" or "faculty focused more on this"?

{previous_context}

## YOUR OUTPUT FORMAT:

Provide a concise response with TWO sections only:

### ✓ STRENGTHS
2-3 bullet points of what the improvement report does well. Acknowledge genuine efforts.

### ✎ REVISIONS REQUESTED
Bullet-pointed list of specific revisions needed. Each bullet should:
- State the issue clearly and briefly
- Be actionable (they should know exactly what to fix)
- Use collegial, supportive language

If no revisions needed, say "No revisions needed - report adequately documents improvement efforts."

Keep your total response under 400 words. Be direct and supportive.

## Report to Analyze:

{report_text}
"""

DEFAULT_PLAN_ANALYSIS_PROMPT = """You are an expert in higher education assessment, supporting colleagues in planning their assessment activities.

{rubric_guidance}

{tone_instructions}

## Your Task: Analyze this Next Cycle Plan

This is a forward-looking plan (no results yet). Run ALL checks internally and provide a concise, actionable response.

### INTERNAL CHECKS TO RUN (do not output these sections):

1. **Field Completeness** - All required fields filled (Outcome Rationale is optional)

2. **Competencies/Functions** - Must be coherent statements (not single words), and each outcome's Related Competency/Function must match one listed in General Information

3. **Bloom's Taxonomy** - Action verbs appropriate for program level:
   - UG lower: Levels 1-3 OK
   - UG upper/capstone: Levels 3-5 expected
   - Graduate: Levels 4-6 expected (1-3 too low)
   - Doctoral: Levels 5-6 expected
   - Flag vague verbs: "understand," "know," "learn," "appreciate"

4. **Outcome Labels**:
   - Academic programs: "Student Learning Outcomes"
   - Administrative units: "Outcomes" (NOT "learning outcomes")

5. **Methodology** - Will it produce quantitative results? Is it appropriate for the outcome?

6. **Criteria for Success** - Clear, specific, appropriately ambitious

7. **Action Steps** - Specific actions to facilitate achievement (who, what, when), not vague or identical across outcomes

8. **Alignment** - Strategic plan theme and competency/function alignment

{custom_rubric}

## YOUR OUTPUT FORMAT:

Provide a concise response with TWO sections only:

### ✓ STRENGTHS
2-3 bullet points of what the plan does well. Be genuine and specific.

### ✎ REVISIONS REQUESTED
Bullet-pointed list of specific revisions needed. Each bullet should:
- State the issue clearly and briefly
- Be actionable (they should know exactly what to fix)
- Use collegial, supportive language

If no revisions needed, say "No revisions needed - plan meets all criteria."

Keep your total response under 500 words. Be direct and actionable.

## Plan to Analyze:

{report_text}
"""

# ============================================================================
# SESSION STATE INITIALIZATION
# ============================================================================

def init_session_state():
    """Initialize all session state variables."""
    defaults = {
        "authenticated": False,
        "admin_mode": False,
        "rubric_guidance": DEFAULT_RUBRIC_GUIDANCE,
        "tone_instructions": DEFAULT_TONE_INSTRUCTIONS,
        "results_prompt": DEFAULT_RESULTS_ANALYSIS_PROMPT,
        "improvement_prompt": DEFAULT_IMPROVEMENT_ANALYSIS_PROMPT,
        "plan_prompt": DEFAULT_PLAN_ANALYSIS_PROMPT,
        "custom_rubric_text": "",
        "good_outcome_example": "",
        "good_criteria_example": "",
        "good_improvement_example": "",
        "good_action_example": "",
        "results": None,
        "extracted_metadata": None,
        "filename": None,
        "unit_registry": None,
        "gsheet_connected": False,
        "batch_mode": False,
        "batch_files": [],
        "batch_metadata": [],
        "copy_text": None
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_session_state()

# ============================================================================
# AUTHENTICATION
# ============================================================================

def check_password():
    """Password protection with admin mode option."""
    
    def password_entered():
        user_pass = st.session_state.get("password", "")
        admin_pass = os.environ.get("ADMIN_PASSWORD", "admin_assessment2024")
        regular_pass = os.environ.get("APP_PASSWORD", "assessment2024")
        
        if user_pass == admin_pass:
            st.session_state["authenticated"] = True
            st.session_state["admin_mode"] = True
        elif user_pass == regular_pass:
            st.session_state["authenticated"] = True
            st.session_state["admin_mode"] = False
        else:
            st.session_state["authenticated"] = False
        
        if "password" in st.session_state:
            del st.session_state["password"]

    if not st.session_state["authenticated"]:
        # Inject branding for login page
        inject_uta_branding()
        
        # Clean login page
        st.markdown("""
        <div style="text-align: center; padding: 3rem 0 2rem 0;">
            <h1 style="color: #003865; font-weight: 700; margin-bottom: 0.5rem;">Assessment Analyzer</h1>
            <p style="color: #5a6777; font-size: 1rem;">Office of Institutional Effectiveness and Reporting</p>
            <p style="color: #8896a6; font-size: 0.875rem;">The University of Texas at Arlington</p>
            <div style="width: 60px; height: 3px; background-color: #F58025; margin: 1.5rem auto 0 auto;"></div>
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns([1, 1.5, 1])
        with col2:
            st.text_input(
                "Password", 
                type="password", 
                key="password",
                on_change=password_entered,
                placeholder="Enter your password"
            )
            
            st.markdown("<br>", unsafe_allow_html=True)
            st.caption("Contact your administrator for access credentials.")
        
        return False
    
    return True

# ============================================================================
# FILE PROCESSING
# ============================================================================

def extract_text_from_pdf(file) -> str:
    """Extract text from PDF file."""
    if not PDF_SUPPORT:
        st.error("PDF support not available. Please install PyPDF2.")
        return ""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""

def extract_text_from_docx(file) -> str:
    """Extract text from Word document."""
    if not DOCX_SUPPORT:
        st.error("Word document support not available. Please install python-docx.")
        return ""
    try:
        doc = Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " | "
                text += "\n"
        return text
    except Exception as e:
        st.error(f"Error reading Word document: {str(e)}")
        return ""

def process_uploaded_file(uploaded_file) -> str:
    """Process uploaded file and extract text."""
    if uploaded_file is None:
        return ""
    
    file_type = uploaded_file.type
    file_name = uploaded_file.name.lower()
    
    if file_type == "application/pdf" or file_name.endswith('.pdf'):
        return extract_text_from_pdf(uploaded_file)
    elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"] or file_name.endswith('.docx'):
        return extract_text_from_docx(uploaded_file)
    elif file_type == "text/plain" or file_name.endswith('.txt'):
        return uploaded_file.read().decode("utf-8")
    else:
        st.error(f"Unsupported file type: {file_type}")
        return ""

# ============================================================================
# UNIT REGISTRY MANAGEMENT
# ============================================================================

def load_unit_registry():
    """Load unit registry from session or initialize from files."""
    if st.session_state.get("unit_registry") is not None:
        return st.session_state["unit_registry"]
    
    # Try to load from bundled CSV files
    registry = {"academic": [], "administrative": []}
    
    # Check for local files first
    academic_path = "unit_registry_academic.csv"
    admin_path = "unit_registry_admin.csv"
    
    try:
        if os.path.exists(academic_path):
            df = pd.read_csv(academic_path)
            registry["academic"] = df.to_dict('records')
    except Exception as e:
        st.warning(f"Could not load academic registry: {e}")
    
    try:
        if os.path.exists(admin_path):
            df = pd.read_csv(admin_path)
            registry["administrative"] = df.to_dict('records')
    except Exception as e:
        st.warning(f"Could not load administrative registry: {e}")
    
    st.session_state["unit_registry"] = registry
    return registry

def generate_unit_id(unit_name: str, college_dept: str, unit_type: str) -> str:
    """Generate a unique unit ID."""
    # Create a short hash for uniqueness
    combined = f"{college_dept}_{unit_name}_{unit_type}"
    short_hash = hashlib.md5(combined.encode()).hexdigest()[:6]
    # Clean the college_dept for the ID
    clean_dept = re.sub(r'[^A-Za-z0-9]', '', college_dept)[:10]
    return f"{clean_dept}-{short_hash}".upper()

def find_matching_unit(extracted_name: str, unit_type: str, registry: dict) -> dict:
    """Find matching unit in registry, checking canonical and previous names."""
    registry_list = registry.get("academic" if unit_type == "Academic" else "administrative", [])
    
    extracted_lower = extracted_name.lower().strip()
    
    # First pass: exact match on canonical name
    for unit in registry_list:
        canonical = unit.get("canonical_name", "").lower().strip()
        if canonical == extracted_lower:
            return {"match": unit, "match_type": "exact", "confidence": "high"}
    
    # Second pass: check previous names
    for unit in registry_list:
        previous = unit.get("previous_names", "")
        if previous and isinstance(previous, str):
            prev_list = [p.strip().lower() for p in previous.split(";")]
            if extracted_lower in prev_list:
                return {"match": unit, "match_type": "previous_name", "confidence": "high"}
    
    # Third pass: fuzzy matching (contains)
    for unit in registry_list:
        canonical = unit.get("canonical_name", "").lower().strip()
        # Check if one contains the other (handles slight variations)
        if extracted_lower in canonical or canonical in extracted_lower:
            return {"match": unit, "match_type": "fuzzy", "confidence": "medium"}
    
    # Fourth pass: key terms matching
    extracted_terms = set(re.findall(r'\b\w+\b', extracted_lower))
    best_match = None
    best_score = 0
    
    for unit in registry_list:
        canonical = unit.get("canonical_name", "").lower()
        canonical_terms = set(re.findall(r'\b\w+\b', canonical))
        
        # Calculate overlap
        common = extracted_terms & canonical_terms
        if len(common) > best_score and len(common) >= 2:
            best_score = len(common)
            best_match = unit
    
    if best_match and best_score >= 3:
        return {"match": best_match, "match_type": "terms", "confidence": "low"}
    
    return {"match": None, "match_type": "none", "confidence": "none"}

# ============================================================================
# EXCEL ONLINE (MICROSOFT GRAPH API) INTEGRATION
# ============================================================================

def get_graph_access_token(client_id: str, client_secret: str, tenant_id: str) -> str:
    """Get Microsoft Graph API access token using client credentials."""
    if not EXCEL_ONLINE_SUPPORT:
        return None
    
    try:
        authority = f"https://login.microsoftonline.com/{tenant_id}"
        app = msal.ConfidentialClientApplication(
            client_id,
            authority=authority,
            client_credential=client_secret
        )
        
        # Get token for Microsoft Graph
        result = app.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])
        
        if "access_token" in result:
            return result["access_token"]
        else:
            st.error(f"Failed to get access token: {result.get('error_description', 'Unknown error')}")
            return None
            
    except Exception as e:
        st.error(f"Authentication error: {str(e)}")
        return None

def get_excel_workbook_info(access_token: str, site_id: str, file_path: str) -> dict:
    """Get workbook information from SharePoint/OneDrive."""
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }
    
    # If site_id is provided, use SharePoint; otherwise use OneDrive
    if site_id:
        # SharePoint path
        base_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drive/root:/{file_path}"
    else:
        # OneDrive path (for the app's service account)
        base_url = f"https://graph.microsoft.com/v1.0/drive/root:/{file_path}"
    
    try:
        response = requests.get(base_url, headers=headers)
        if response.status_code == 200:
            return response.json()
        else:
            return {"error": response.text}
    except Exception as e:
        return {"error": str(e)}

def get_or_create_worksheet(access_token: str, drive_id: str, item_id: str, sheet_name: str, headers_list: list) -> bool:
    """Get existing worksheet or create with headers."""
    api_headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }
    
    base_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}/workbook/worksheets"
    
    try:
        # Check if worksheet exists
        response = requests.get(base_url, headers=api_headers)
        if response.status_code != 200:
            return False
        
        worksheets = response.json().get("value", [])
        sheet_exists = any(ws.get("name") == sheet_name for ws in worksheets)
        
        if not sheet_exists:
            # Create worksheet
            create_response = requests.post(
                base_url,
                headers=api_headers,
                json={"name": sheet_name}
            )
            if create_response.status_code not in [200, 201]:
                st.error(f"Failed to create worksheet: {create_response.text}")
                return False
            
            # Add headers to new worksheet
            header_range = f"A1:{chr(65 + len(headers_list) - 1)}1"
            range_url = f"{base_url}/{sheet_name}/range(address='{header_range}')"
            
            update_response = requests.patch(
                range_url,
                headers=api_headers,
                json={"values": [headers_list]}
            )
            if update_response.status_code not in [200, 201]:
                st.warning(f"Created worksheet but failed to add headers: {update_response.text}")
        
        return True
        
    except Exception as e:
        st.error(f"Error managing worksheet: {str(e)}")
        return False

def get_worksheet_data(access_token: str, drive_id: str, item_id: str, sheet_name: str) -> list:
    """Get all data from a worksheet."""
    api_headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }
    
    url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}/workbook/worksheets/{sheet_name}/usedRange"
    
    try:
        response = requests.get(url, headers=api_headers)
        if response.status_code == 200:
            data = response.json()
            values = data.get("values", [])
            if len(values) > 1:
                headers = values[0]
                records = []
                for row in values[1:]:
                    record = {}
                    for i, header in enumerate(headers):
                        record[header] = row[i] if i < len(row) else ""
                    records.append(record)
                return records
            return []
        else:
            return []
    except Exception as e:
        return []

def save_metadata_to_excel_online(access_token: str, drive_id: str, item_id: str, 
                                   metadata_rows: list, report_type: str) -> bool:
    """Save metadata to appropriate worksheet in Excel Online, handling duplicates."""
    
    api_headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }
    
    # Determine sheet name and headers based on report type
    if report_type == "Results Report":
        sheet_name = "Results_Data"
        headers = [
            "unit_id", "unit_type", "unit_name", "college_division", "degree_level",
            "modality", "academic_year", "report_type", "outcome_id", "outcome_text",
            "outcome_label", "related_competency_or_function", "strategic_plan_theme",
            "core_objective", "assessment_method", "assessment_method_normalized",
            "sample_size", "benchmark", "result_value", "achievement_level",
            "gap_from_benchmark", "proposed_improvement", "responsible_party",
            "improvement_timeline", "upload_timestamp", "last_updated"
        ]
    elif report_type == "Improvement Report":
        sheet_name = "Improvement_Data"
        headers = [
            "unit_id", "unit_type", "unit_name", "college_division",
            "academic_year", "report_type", "outcome_id", "improvement_action_taken",
            "connection_to_previous", "previous_proposal_text",
            "upload_timestamp", "last_updated"
        ]
    else:  # Next Cycle Plan
        sheet_name = "Plan_Data"
        headers = [
            "unit_id", "unit_type", "unit_name", "college_division", "degree_level",
            "academic_year", "report_type", "outcome_id", "outcome_text",
            "outcome_label", "related_competency_or_function", "strategic_plan_theme",
            "core_objective", "planned_method", "planned_benchmark",
            "action_steps", "responsible_party", "upload_timestamp", "last_updated"
        ]
    
    # Ensure worksheet exists
    if not get_or_create_worksheet(access_token, drive_id, item_id, sheet_name, headers):
        return False
    
    try:
        # Get existing data
        existing_data = get_worksheet_data(access_token, drive_id, item_id, sheet_name)
        
        # Build index of existing rows by unique key
        existing_index = {}
        for i, row in enumerate(existing_data):
            key = f"{row.get('unit_id', '')}|{row.get('academic_year', '')}|{row.get('outcome_id', '')}"
            existing_index[key] = i + 2  # +2 for header row and 1-based indexing
        
        # Track keys for orphan detection
        new_keys = set()
        unit_year_combos = set()
        
        rows_to_update = []
        rows_to_append = []
        
        for row in metadata_rows:
            key = f"{row.get('unit_id', '')}|{row.get('academic_year', '')}|{row.get('outcome_id', '')}"
            new_keys.add(key)
            unit_year_combos.add(f"{row.get('unit_id', '')}|{row.get('academic_year', '')}")
            
            # Prepare row data in header order
            row_data = [str(row.get(h, "")) for h in headers]
            
            if key in existing_index:
                rows_to_update.append({
                    "row_num": existing_index[key],
                    "data": row_data
                })
            else:
                rows_to_append.append(row_data)
        
        base_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}/workbook/worksheets/{sheet_name}"
        
        # Update existing rows
        for update in rows_to_update:
            row_num = update["row_num"]
            range_address = f"A{row_num}:{chr(65 + len(headers) - 1)}{row_num}"
            range_url = f"{base_url}/range(address='{range_address}')"
            
            response = requests.patch(
                range_url,
                headers=api_headers,
                json={"values": [update["data"]]}
            )
            if response.status_code not in [200, 201]:
                st.warning(f"Failed to update row {row_num}")
        
        # Append new rows
        if rows_to_append:
            # Find the next empty row
            next_row = len(existing_data) + 2  # +1 for header, +1 for 1-based
            
            for i, row_data in enumerate(rows_to_append):
                row_num = next_row + i
                range_address = f"A{row_num}:{chr(65 + len(headers) - 1)}{row_num}"
                range_url = f"{base_url}/range(address='{range_address}')"
                
                response = requests.patch(
                    range_url,
                    headers=api_headers,
                    json={"values": [row_data]}
                )
        
        # Delete orphaned outcomes (same unit+year but outcome not in new data)
        rows_to_delete = []
        for key, row_num in existing_index.items():
            parts = key.split("|")
            if len(parts) >= 2:
                unit_year = f"{parts[0]}|{parts[1]}"
                if unit_year in unit_year_combos and key not in new_keys:
                    rows_to_delete.append(row_num)
        
        # Delete from bottom up (to preserve row numbers)
        for row_num in sorted(rows_to_delete, reverse=True):
            # Clear the row (Excel Online doesn't have direct row delete via Graph API easily)
            range_address = f"A{row_num}:{chr(65 + len(headers) - 1)}{row_num}"
            range_url = f"{base_url}/range(address='{range_address}')"
            
            # Clear the row content
            requests.patch(
                range_url,
                headers=api_headers,
                json={"values": [[""] * len(headers)]}
            )
        
        return True
        
    except Exception as e:
        st.error(f"Error saving to Excel Online: {str(e)}")
        return False

def get_historical_data_excel(access_token: str, drive_id: str, item_id: str, 
                               unit_id: str, outcome_id: str = None) -> list:
    """Retrieve historical data for stagnation detection and context."""
    try:
        data = get_worksheet_data(access_token, drive_id, item_id, "Results_Data")
        
        # Filter by unit_id
        filtered = [r for r in data if r.get("unit_id") == unit_id]
        
        # If outcome_id specified, filter further
        if outcome_id:
            filtered = [r for r in filtered if r.get("outcome_id") == outcome_id]
        
        return filtered
        
    except Exception as e:
        return []

def get_previous_improvements_excel(access_token: str, drive_id: str, item_id: str, unit_id: str) -> list:
    """Get proposed improvements from previous Results reports."""
    historical = get_historical_data_excel(access_token, drive_id, item_id, unit_id)
    
    # Sort by academic year descending
    historical.sort(key=lambda x: x.get("academic_year", ""), reverse=True)
    
    improvements = []
    for record in historical:
        if record.get("proposed_improvement"):
            improvements.append({
                "academic_year": record.get("academic_year"),
                "outcome_id": record.get("outcome_id"),
                "proposed_improvement": record.get("proposed_improvement")
            })
    
    return improvements

# ============================================================================
# STAGNATION DETECTION
# ============================================================================

def check_stagnation(access_token: str, drive_id: str, item_id: str, unit_id: str, 
                     outcome_id: str, current_method: str, current_year: str) -> dict:
    """Check if outcome has been achieved with same methodology for 3+ years."""
    
    historical = get_historical_data_excel(access_token, drive_id, item_id, unit_id, outcome_id)
    
    if len(historical) < 2:  # Need at least 2 previous years
        return {"stagnant": False, "reason": "insufficient_history"}
    
    # Sort by academic year
    historical.sort(key=lambda x: x.get("academic_year", ""))
    
    # Get last 3 years including current
    recent = historical[-3:] if len(historical) >= 3 else historical
    
    # Check if all achieved
    all_achieved = all(
        r.get("achievement_level") == "Fully Achieved" 
        for r in recent
    )
    
    if not all_achieved:
        return {"stagnant": False, "reason": "not_all_achieved"}
    
    # Check methodology similarity (will be evaluated by AI in analysis)
    methods = [r.get("assessment_method_normalized", r.get("assessment_method", "")) for r in recent]
    
    return {
        "stagnant": True,
        "years_achieved": len(recent),
        "years": [r.get("academic_year") for r in recent],
        "methods": methods,
        "needs_ai_verification": True
    }

# ============================================================================
# METADATA EXTRACTION
# ============================================================================

def extract_metadata_with_ai(report_text: str, report_type: str, api_key: str) -> dict:
    """Use Claude to extract structured metadata from report."""
    
    client = anthropic.Anthropic(api_key=api_key)
    
    if report_type == "Results Report":
        extraction_prompt = """Extract metadata from this assessment report. Return ONLY valid JSON.

{
  "unit_type": "Academic or Administrative (check header of report)",
  "unit_name": "Full program/unit name",
  "college_division": "College or Division code/name",
  "degree_level": "UG/GR/Doctoral/Certificate/NA",
  "modality": "On-campus/Online/Hybrid/Multiple or specific description",
  "academic_year": "YYYY-YYYY format",
  "outcomes": [
    {
      "outcome_id": "Short identifier or name",
      "outcome_text": "Full outcome statement",
      "related_competency_or_function": "Related Student Competency or Core Function as stated",
      "strategic_plan_theme": "If mapped to strategic plan",
      "core_objective": "If mapped to core objectives",
      "assessment_method": "Full description of how assessed",
      "assessment_method_normalized": "Simplified 5-10 word summary for comparison",
      "sample_size": "Number or description",
      "benchmark": "Exact criteria for success as stated",
      "result_value": "Exact results as stated",
      "achievement_level": "Fully Achieved/Partially Achieved/Not Achieved/Inconclusive",
      "gap_from_benchmark": "Numeric if calculable, otherwise description",
      "proposed_improvement": "What they plan to do",
      "responsible_party": "Who is responsible",
      "improvement_timeline": "When"
    }
  ]
}

Report to extract from:
"""
    elif report_type == "Improvement Report":
        extraction_prompt = """Extract metadata from this improvement report. Return ONLY valid JSON.

{
  "unit_type": "Academic or Administrative",
  "unit_name": "Full program/unit name",
  "college_division": "College or Division",
  "academic_year": "YYYY-YYYY format (year improvements were implemented)",
  "improvements": [
    {
      "outcome_id": "What outcome/area this addresses",
      "improvement_action_taken": "Detailed description of what was done",
      "connection_to_previous": "Yes/No/Partial - does it reference original finding?"
    }
  ]
}

Report to extract from:
"""
    else:  # Next Cycle Plan
        extraction_prompt = """Extract metadata from this assessment plan. Return ONLY valid JSON.

{
  "unit_type": "Academic or Administrative",
  "unit_name": "Full program/unit name", 
  "college_division": "College or Division",
  "degree_level": "UG/GR/Doctoral/Certificate/NA",
  "academic_year": "YYYY-YYYY format (year being planned)",
  "outcomes": [
    {
      "outcome_id": "Short identifier",
      "outcome_text": "Full outcome statement",
      "related_competency_or_function": "Related competency or core function",
      "strategic_plan_theme": "If mapped",
      "core_objective": "If mapped",
      "planned_method": "How they will assess",
      "planned_benchmark": "Proposed criteria for success",
      "action_steps": "How students will be prepared",
      "responsible_party": "Who is responsible"
    }
  ]
}

Report to extract from:
"""
    
    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{
                "role": "user",
                "content": extraction_prompt + report_text
            }]
        )
        
        response_text = response.content[0].text
        
        # Parse JSON from response
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            return json.loads(json_match.group())
        else:
            return {"error": "Could not parse JSON from response"}
            
    except Exception as e:
        return {"error": str(e)}

# ============================================================================
# AI ANALYSIS
# ============================================================================

def analyze_report(report_text: str, report_type: str, api_key: str,
                   stagnation_info: dict = None, previous_improvements: list = None) -> dict:
    """Analyze report using Claude with appropriate prompt."""
    
    client = anthropic.Anthropic(api_key=api_key)
    
    # Build context sections
    stagnation_context = ""
    if stagnation_info and stagnation_info.get("stagnant"):
        stagnation_context = f"""
## STAGNATION CHECK REQUIRED

Historical data shows this outcome may have been Fully Achieved for {stagnation_info['years_achieved']} consecutive years:
Years: {', '.join(stagnation_info['years'])}
Methods used: {'; '.join(stagnation_info['methods'])}

Please evaluate:
1. Are these methodologies functionally the same? (Same test/exam in same course counts as same, even if wording differs. Different test OR different course = different methodology)
2. If methodology is unchanged AND outcome achieved for 3+ years, include a gentle note suggesting the unit consider evolving their assessment (raising benchmark, refining outcome, changing methodology, or replacing with more challenging outcome).

Frame this supportively - sustained achievement is positive, but assessment should evolve over time.
"""
    
    previous_context = ""
    if previous_improvements:
        prev_text = "\n".join([
            f"- {p['academic_year']}, {p['outcome_id']}: {p['proposed_improvement'][:200]}..."
            for p in previous_improvements[:5]
        ])
        previous_context = f"""
## PREVIOUS PROPOSED IMPROVEMENTS

The following improvements were proposed in previous assessment cycles:
{prev_text}

When analyzing this improvement report, consider whether the actions taken align with what was previously proposed. Note alignments positively, and gently note if proposed improvements don't appear to have been addressed.
"""
    
    custom_rubric = ""
    if st.session_state.get("custom_rubric_text"):
        custom_rubric = f"""
## ADDITIONAL EVALUATION CRITERIA (Custom Rubric)

{st.session_state['custom_rubric_text']}
"""
    
    # Build good examples section
    good_examples = ""
    examples_list = []
    
    if st.session_state.get("good_outcome_example"):
        examples_list.append(f"**Good Outcome Statement:**\n{st.session_state['good_outcome_example']}")
    
    if st.session_state.get("good_criteria_example"):
        examples_list.append(f"**Good Criteria for Success:**\n{st.session_state['good_criteria_example']}")
    
    if st.session_state.get("good_improvement_example"):
        examples_list.append(f"**Good Proposed Improvement:**\n{st.session_state['good_improvement_example']}")
    
    if st.session_state.get("good_action_example"):
        examples_list.append(f"**Good Action Steps:**\n{st.session_state['good_action_example']}")
    
    if examples_list:
        good_examples = "\n\n## EXAMPLES OF QUALITY WORK\nUse these as reference for what good looks like. Recognize similar quality in reports you analyze:\n\n" + "\n\n".join(examples_list)
    
    # Combine custom rubric and examples
    custom_rubric = custom_rubric + good_examples
    
    # Select appropriate prompt template
    if report_type == "Results Report":
        prompt_template = st.session_state.get("results_prompt", DEFAULT_RESULTS_ANALYSIS_PROMPT)
    elif report_type == "Improvement Report":
        prompt_template = st.session_state.get("improvement_prompt", DEFAULT_IMPROVEMENT_ANALYSIS_PROMPT)
    else:
        prompt_template = st.session_state.get("plan_prompt", DEFAULT_PLAN_ANALYSIS_PROMPT)
    
    # Build final prompt
    prompt = prompt_template.format(
        rubric_guidance=st.session_state.get("rubric_guidance", DEFAULT_RUBRIC_GUIDANCE),
        tone_instructions=st.session_state.get("tone_instructions", DEFAULT_TONE_INSTRUCTIONS),
        stagnation_context=stagnation_context,
        previous_context=previous_context,
        custom_rubric=custom_rubric,
        report_text=report_text
    )
    
    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{"role": "user", "content": prompt}]
        )
        
        analysis_text = response.content[0].text
        
        # Calculate costs
        input_tokens = response.usage.input_tokens
        output_tokens = response.usage.output_tokens
        estimated_cost = (input_tokens * 0.003 / 1000) + (output_tokens * 0.015 / 1000)
        
        return {
            "analysis": analysis_text,
            "tokens": {"input": input_tokens, "output": output_tokens},
            "cost": f"${estimated_cost:.4f}"
        }
        
    except Exception as e:
        return {"error": str(e)}

# ============================================================================
# METADATA PREVIEW AND EDITING UI
# ============================================================================

def render_metadata_editor(metadata: dict, report_type: str, registry: dict):
    """Render editable metadata preview."""
    
    st.subheader("📝 Review and Edit Metadata")
    st.caption("Correct any extraction errors before saving to Google Sheets.")
    
    edited = metadata.copy()
    
    # Unit-level fields
    col1, col2 = st.columns(2)
    
    with col1:
        edited["unit_type"] = st.selectbox(
            "Unit Type",
            ["Academic", "Administrative"],
            index=0 if metadata.get("unit_type", "").lower() == "academic" else 1,
            key="edit_unit_type"
        )
        
        edited["unit_name"] = st.text_input(
            "Unit Name",
            value=metadata.get("unit_name", ""),
            key="edit_unit_name"
        )
        
        # Unit matching
        match_result = find_matching_unit(
            edited["unit_name"], 
            edited["unit_type"],
            registry
        )
        
        if match_result["match"]:
            if match_result["confidence"] == "high":
                st.success(f"✓ Matched: {match_result['match'].get('canonical_name')}")
                edited["unit_id"] = match_result["match"].get("unit_id", "")
                edited["canonical_name"] = match_result["match"].get("canonical_name", "")
            else:
                st.warning(f"⚠ Possible match: {match_result['match'].get('canonical_name')}")
                if st.checkbox("Confirm this match", key="confirm_match"):
                    edited["unit_id"] = match_result["match"].get("unit_id", "")
                    edited["canonical_name"] = match_result["match"].get("canonical_name", "")
                else:
                    edited["unit_id"] = generate_unit_id(
                        edited["unit_name"],
                        edited.get("college_division", ""),
                        edited["unit_type"]
                    )
        else:
            st.info("ℹ No existing match found. Will create new registry entry.")
            edited["unit_id"] = generate_unit_id(
                edited["unit_name"],
                edited.get("college_division", ""),
                edited["unit_type"]
            )
    
    with col2:
        edited["college_division"] = st.text_input(
            "College/Division",
            value=metadata.get("college_division", ""),
            key="edit_college"
        )
        
        edited["academic_year"] = st.text_input(
            "Academic Year",
            value=metadata.get("academic_year", ""),
            key="edit_year"
        )
        
        if report_type != "Improvement Report":
            edited["degree_level"] = st.selectbox(
                "Degree Level",
                ["UG", "GR", "Doctoral", "Certificate", "NA"],
                index=["UG", "GR", "Doctoral", "Certificate", "NA"].index(
                    metadata.get("degree_level", "NA")
                ) if metadata.get("degree_level") in ["UG", "GR", "Doctoral", "Certificate", "NA"] else 4,
                key="edit_degree"
            )
    
    # Outcomes
    st.divider()
    outcomes_key = "outcomes" if report_type != "Improvement Report" else "improvements"
    outcomes = metadata.get(outcomes_key, [])
    
    edited_outcomes = []
    
    for i, outcome in enumerate(outcomes):
        with st.expander(f"{'Outcome' if report_type != 'Improvement Report' else 'Improvement'} {i+1}: {outcome.get('outcome_id', outcome.get('outcome_id', 'Untitled'))[:50]}...", expanded=(i==0)):
            edited_outcome = outcome.copy()
            
            if report_type == "Results Report":
                col1, col2 = st.columns(2)
                with col1:
                    edited_outcome["outcome_id"] = st.text_input(
                        "Outcome ID",
                        value=outcome.get("outcome_id", ""),
                        key=f"outcome_id_{i}"
                    )
                    edited_outcome["outcome_text"] = st.text_area(
                        "Outcome Text",
                        value=outcome.get("outcome_text", ""),
                        key=f"outcome_text_{i}",
                        height=100
                    )
                    edited_outcome["related_competency_or_function"] = st.text_input(
                        "Related Competency/Function",
                        value=outcome.get("related_competency_or_function", ""),
                        key=f"competency_{i}"
                    )
                    edited_outcome["strategic_plan_theme"] = st.selectbox(
                        "Strategic Plan Theme",
                        [""] + STRATEGIC_THEMES,
                        index=(STRATEGIC_THEMES.index(outcome.get("strategic_plan_theme")) + 1) 
                            if outcome.get("strategic_plan_theme") in STRATEGIC_THEMES else 0,
                        key=f"theme_{i}"
                    )
                    edited_outcome["core_objective"] = st.selectbox(
                        "Core Objective (if applicable)",
                        [""] + CORE_OBJECTIVES,
                        index=(CORE_OBJECTIVES.index(outcome.get("core_objective")) + 1)
                            if outcome.get("core_objective") in CORE_OBJECTIVES else 0,
                        key=f"core_obj_{i}"
                    )
                
                with col2:
                    edited_outcome["assessment_method"] = st.text_area(
                        "Assessment Method",
                        value=outcome.get("assessment_method", ""),
                        key=f"method_{i}",
                        height=80
                    )
                    edited_outcome["benchmark"] = st.text_input(
                        "Benchmark/Criteria for Success",
                        value=outcome.get("benchmark", ""),
                        key=f"benchmark_{i}"
                    )
                    edited_outcome["result_value"] = st.text_input(
                        "Result",
                        value=outcome.get("result_value", ""),
                        key=f"result_{i}"
                    )
                    edited_outcome["achievement_level"] = st.selectbox(
                        "Achievement Level",
                        ACHIEVEMENT_LEVELS,
                        index=ACHIEVEMENT_LEVELS.index(outcome.get("achievement_level"))
                            if outcome.get("achievement_level") in ACHIEVEMENT_LEVELS else 2,
                        key=f"achieved_{i}"
                    )
                    edited_outcome["proposed_improvement"] = st.text_area(
                        "Proposed Improvement",
                        value=outcome.get("proposed_improvement", ""),
                        key=f"improvement_{i}",
                        height=80
                    )
                    edited_outcome["responsible_party"] = st.text_input(
                        "Responsible Party",
                        value=outcome.get("responsible_party", ""),
                        key=f"responsible_{i}"
                    )
            
            elif report_type == "Improvement Report":
                edited_outcome["outcome_id"] = st.text_input(
                    "Outcome/Area Addressed",
                    value=outcome.get("outcome_id", ""),
                    key=f"outcome_id_{i}"
                )
                edited_outcome["improvement_action_taken"] = st.text_area(
                    "Improvement Actions Taken",
                    value=outcome.get("improvement_action_taken", ""),
                    key=f"action_{i}",
                    height=150
                )
                edited_outcome["connection_to_previous"] = st.selectbox(
                    "Connection to Previous Findings",
                    ["Yes", "No", "Partial"],
                    index=["Yes", "No", "Partial"].index(outcome.get("connection_to_previous", "No"))
                        if outcome.get("connection_to_previous") in ["Yes", "No", "Partial"] else 1,
                    key=f"connection_{i}"
                )
            
            else:  # Next Cycle Plan
                col1, col2 = st.columns(2)
                with col1:
                    edited_outcome["outcome_id"] = st.text_input(
                        "Outcome ID",
                        value=outcome.get("outcome_id", ""),
                        key=f"outcome_id_{i}"
                    )
                    edited_outcome["outcome_text"] = st.text_area(
                        "Outcome Text",
                        value=outcome.get("outcome_text", ""),
                        key=f"outcome_text_{i}",
                        height=100
                    )
                    edited_outcome["related_competency_or_function"] = st.text_input(
                        "Related Competency/Function",
                        value=outcome.get("related_competency_or_function", ""),
                        key=f"competency_{i}"
                    )
                    edited_outcome["strategic_plan_theme"] = st.selectbox(
                        "Strategic Plan Theme",
                        [""] + STRATEGIC_THEMES,
                        index=(STRATEGIC_THEMES.index(outcome.get("strategic_plan_theme")) + 1)
                            if outcome.get("strategic_plan_theme") in STRATEGIC_THEMES else 0,
                        key=f"theme_{i}"
                    )
                
                with col2:
                    edited_outcome["planned_method"] = st.text_area(
                        "Planned Assessment Method",
                        value=outcome.get("planned_method", ""),
                        key=f"method_{i}",
                        height=80
                    )
                    edited_outcome["planned_benchmark"] = st.text_input(
                        "Planned Benchmark",
                        value=outcome.get("planned_benchmark", ""),
                        key=f"benchmark_{i}"
                    )
                    edited_outcome["action_steps"] = st.text_area(
                        "Action Steps",
                        value=outcome.get("action_steps", ""),
                        key=f"actions_{i}",
                        height=80
                    )
                    edited_outcome["responsible_party"] = st.text_input(
                        "Responsible Party",
                        value=outcome.get("responsible_party", ""),
                        key=f"responsible_{i}"
                    )
            
            edited_outcomes.append(edited_outcome)
    
    edited[outcomes_key] = edited_outcomes
    
    return edited

def prepare_rows_for_sheet(metadata: dict, report_type: str) -> list:
    """Convert edited metadata to rows for Google Sheets."""
    
    now = datetime.now().isoformat()
    rows = []
    
    base_data = {
        "unit_id": metadata.get("unit_id", ""),
        "unit_type": metadata.get("unit_type", ""),
        "unit_name": metadata.get("canonical_name", metadata.get("unit_name", "")),
        "college_division": metadata.get("college_division", ""),
        "upload_timestamp": now,
        "last_updated": now
    }
    
    if report_type == "Results Report":
        base_data["degree_level"] = metadata.get("degree_level", "")
        base_data["modality"] = metadata.get("modality", "")
        base_data["academic_year"] = metadata.get("academic_year", "")
        base_data["report_type"] = "Results"
        
        for outcome in metadata.get("outcomes", []):
            row = base_data.copy()
            row["outcome_id"] = outcome.get("outcome_id", "")
            row["outcome_text"] = outcome.get("outcome_text", "")[:500]
            row["outcome_label"] = "Student Learning Outcome" if metadata.get("unit_type") == "Academic" else "Outcome"
            row["related_competency_or_function"] = outcome.get("related_competency_or_function", "")
            row["strategic_plan_theme"] = outcome.get("strategic_plan_theme", "")
            row["core_objective"] = outcome.get("core_objective", "")
            row["assessment_method"] = outcome.get("assessment_method", "")
            row["assessment_method_normalized"] = outcome.get("assessment_method_normalized", "")
            row["sample_size"] = outcome.get("sample_size", "")
            row["benchmark"] = outcome.get("benchmark", "")
            row["result_value"] = outcome.get("result_value", "")
            row["achievement_level"] = outcome.get("achievement_level", "")
            row["gap_from_benchmark"] = outcome.get("gap_from_benchmark", "")
            row["proposed_improvement"] = outcome.get("proposed_improvement", "")[:500]
            row["responsible_party"] = outcome.get("responsible_party", "")
            row["improvement_timeline"] = outcome.get("improvement_timeline", "")
            rows.append(row)
    
    elif report_type == "Improvement Report":
        base_data["academic_year"] = metadata.get("academic_year", "")
        base_data["report_type"] = "Improvement"
        
        for improvement in metadata.get("improvements", []):
            row = base_data.copy()
            row["outcome_id"] = improvement.get("outcome_id", "")
            row["improvement_action_taken"] = improvement.get("improvement_action_taken", "")[:1000]
            row["connection_to_previous"] = improvement.get("connection_to_previous", "")
            row["previous_proposal_text"] = improvement.get("previous_proposal_text", "")
            rows.append(row)
    
    else:  # Next Cycle Plan
        base_data["degree_level"] = metadata.get("degree_level", "")
        base_data["academic_year"] = metadata.get("academic_year", "")
        base_data["report_type"] = "Plan"
        
        for outcome in metadata.get("outcomes", []):
            row = base_data.copy()
            row["outcome_id"] = outcome.get("outcome_id", "")
            row["outcome_text"] = outcome.get("outcome_text", "")[:500]
            row["outcome_label"] = "Student Learning Outcome" if metadata.get("unit_type") == "Academic" else "Outcome"
            row["related_competency_or_function"] = outcome.get("related_competency_or_function", "")
            row["strategic_plan_theme"] = outcome.get("strategic_plan_theme", "")
            row["core_objective"] = outcome.get("core_objective", "")
            row["planned_method"] = outcome.get("planned_method", "")
            row["planned_benchmark"] = outcome.get("planned_benchmark", "")
            row["action_steps"] = outcome.get("action_steps", "")[:500]
            row["responsible_party"] = outcome.get("responsible_party", "")
            rows.append(row)
    
    return rows

# ============================================================================
# BATCH IMPORT MODE
# ============================================================================

def render_batch_import(api_key: str, access_token: str, drive_id: str, item_id: str, 
                        excel_connected: bool, registry: dict):
    """Render batch import interface."""
    
    render_uta_header("Batch Import")
    
    st.info("Upload multiple historical reports to populate the metadata database. No analysis will be performed.")
    
    report_type = st.selectbox(
        "Report Type (all files must be same type)",
        REPORT_TYPES,
        key="batch_report_type"
    )
    
    uploaded_files = st.file_uploader(
        "Upload reports",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        key="batch_files"
    )
    
    if uploaded_files:
        st.write(f"**{len(uploaded_files)} files selected**")
        
        if st.button("Extract Metadata from All Files", type="primary"):
            all_metadata = []
            progress = st.progress(0)
            
            for i, file in enumerate(uploaded_files):
                with st.spinner(f"Processing {file.name}..."):
                    text = process_uploaded_file(file)
                    if text:
                        metadata = extract_metadata_with_ai(text, report_type, api_key)
                        if "error" not in metadata:
                            metadata["_filename"] = file.name
                            all_metadata.append(metadata)
                        else:
                            st.warning(f"Error extracting from {file.name}: {metadata['error']}")
                    else:
                        st.warning(f"Could not extract text from {file.name}")
                
                progress.progress((i + 1) / len(uploaded_files))
            
            st.session_state["batch_metadata"] = all_metadata
            st.success(f"Extracted metadata from {len(all_metadata)} files")
    
    # Show extracted metadata for review
    if st.session_state.get("batch_metadata"):
        st.divider()
        st.subheader("Review Extracted Metadata")
        
        for i, metadata in enumerate(st.session_state["batch_metadata"]):
            with st.expander(f"📄 {metadata.get('_filename', f'File {i+1}')} - {metadata.get('unit_name', 'Unknown')}"):
                st.json(metadata)
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("💾 Save All to Excel Online", type="primary"):
                if not excel_connected or not access_token:
                    st.error("Excel Online not connected. Configure in sidebar.")
                else:
                    success_count = 0
                    for metadata in st.session_state["batch_metadata"]:
                        # Generate unit_id
                        match = find_matching_unit(
                            metadata.get("unit_name", ""),
                            metadata.get("unit_type", "Academic"),
                            registry
                        )
                        if match["match"]:
                            metadata["unit_id"] = match["match"].get("unit_id")
                            metadata["canonical_name"] = match["match"].get("canonical_name")
                        else:
                            metadata["unit_id"] = generate_unit_id(
                                metadata.get("unit_name", ""),
                                metadata.get("college_division", ""),
                                metadata.get("unit_type", "Academic")
                            )
                        
                        rows = prepare_rows_for_sheet(metadata, report_type)
                        if save_metadata_to_excel_online(access_token, drive_id, item_id, rows, report_type):
                            success_count += 1
                    
                    st.success(f"Saved {success_count} of {len(st.session_state['batch_metadata'])} reports")
                    st.session_state["batch_metadata"] = []
        
        with col2:
            if st.button("🗑️ Clear All"):
                st.session_state["batch_metadata"] = []
                st.rerun()

# ============================================================================
# ADMIN CONFIGURATION PANEL
# ============================================================================

def render_admin_panel():
    """Render admin configuration interface."""
    
    render_uta_header("Configuration")
    
    st.info("Edit analysis criteria, prompts, and add quality examples. Changes apply to all subsequent analyses.")
    
    tabs = st.tabs([
        "Good Examples",
        "Rubric Guidance",
        "Tone Instructions", 
        "Results Prompt",
        "Improvement Prompt",
        "Plan Prompt",
        "Custom Rubric",
        "Unit Registry"
    ])
    
    # NEW: Good Examples Tab
    with tabs[0]:
        st.subheader("📝 Good Examples")
        st.markdown("""
        Paste examples of **well-written** report sections here. These will be included in the AI's context 
        to help it understand what good looks like. The AI will use these as reference when analyzing reports.
        """)
        
        st.caption("**Example of a Good Outcome Statement:**")
        if "good_outcome_example" not in st.session_state:
            st.session_state["good_outcome_example"] = ""
        st.session_state["good_outcome_example"] = st.text_area(
            "Paste a well-written outcome statement:",
            value=st.session_state.get("good_outcome_example", ""),
            height=100,
            key="example_outcome",
            placeholder="e.g., Students will analyze complex business cases and recommend evidence-based solutions that consider ethical, legal, and economic factors."
        )
        
        st.caption("**Example of Good Criteria for Success:**")
        if "good_criteria_example" not in st.session_state:
            st.session_state["good_criteria_example"] = ""
        st.session_state["good_criteria_example"] = st.text_area(
            "Paste well-written criteria for success:",
            value=st.session_state.get("good_criteria_example", ""),
            height=100,
            key="example_criteria",
            placeholder="e.g., 80% of students (n≥30) will score 3 or higher on the case analysis rubric (4-point scale) as evaluated by two independent faculty raters."
        )
        
        st.caption("**Example of Good Proposed Improvement:**")
        if "good_improvement_example" not in st.session_state:
            st.session_state["good_improvement_example"] = ""
        st.session_state["good_improvement_example"] = st.text_area(
            "Paste a well-written proposed improvement:",
            value=st.session_state.get("good_improvement_example", ""),
            height=100,
            key="example_improvement",
            placeholder="e.g., The MKTG 3310 instructor will integrate two additional case studies focusing on ethical marketing dilemmas during Fall 2025. Assessment will be repeated in Spring 2026 using the same rubric to measure improvement."
        )
        
        st.caption("**Example of Good Action Steps:**")
        if "good_action_example" not in st.session_state:
            st.session_state["good_action_example"] = ""
        st.session_state["good_action_example"] = st.text_area(
            "Paste well-written action steps:",
            value=st.session_state.get("good_action_example", ""),
            height=100,
            key="example_action",
            placeholder="e.g., (1) Dr. Smith will revise ACCT 3310 syllabus to include weekly practice problems by Aug 2025. (2) Department will offer tutoring sessions twice weekly starting Fall 2025. (3) Assessment coordinator will collect and analyze results in Dec 2025."
        )
        
        st.success("These examples will be included when analyzing reports to guide the AI toward recognizing quality work.")
    
    with tabs[1]:
        st.subheader("Rubric Guidance")
        st.caption("Criteria for evaluating reports. Referenced in feedback.")
        st.session_state["rubric_guidance"] = st.text_area(
            "Edit rubric guidance:",
            value=st.session_state["rubric_guidance"],
            height=400,
            key="edit_rubric_guidance"
        )
        if st.button("Reset to Default", key="reset_rubric"):
            st.session_state["rubric_guidance"] = DEFAULT_RUBRIC_GUIDANCE
            st.rerun()
    
    with tabs[2]:
        st.subheader("Tone Instructions")
        st.caption("Controls communication style in feedback.")
        st.session_state["tone_instructions"] = st.text_area(
            "Edit tone instructions:",
            value=st.session_state["tone_instructions"],
            height=400,
            key="edit_tone"
        )
        if st.button("Reset to Default", key="reset_tone"):
            st.session_state["tone_instructions"] = DEFAULT_TONE_INSTRUCTIONS
            st.rerun()
    
    with tabs[3]:
        st.subheader("Results Report Analysis Prompt")
        st.caption("Template for analyzing results reports. The AI runs all checks internally and outputs only Strengths and Revisions Requested.")
        st.session_state["results_prompt"] = st.text_area(
            "Edit prompt:",
            value=st.session_state["results_prompt"],
            height=400,
            key="edit_results_prompt"
        )
        if st.button("Reset to Default", key="reset_results"):
            st.session_state["results_prompt"] = DEFAULT_RESULTS_ANALYSIS_PROMPT
            st.rerun()
    
    with tabs[4]:
        st.subheader("Improvement Report Analysis Prompt")
        st.session_state["improvement_prompt"] = st.text_area(
            "Edit prompt:",
            value=st.session_state["improvement_prompt"],
            height=400,
            key="edit_improvement_prompt"
        )
        if st.button("Reset to Default", key="reset_improvement"):
            st.session_state["improvement_prompt"] = DEFAULT_IMPROVEMENT_ANALYSIS_PROMPT
            st.rerun()
    
    with tabs[5]:
        st.subheader("Next Cycle Plan Analysis Prompt")
        st.session_state["plan_prompt"] = st.text_area(
            "Edit prompt:",
            value=st.session_state["plan_prompt"],
            height=400,
            key="edit_plan_prompt"
        )
        if st.button("Reset to Default", key="reset_plan"):
            st.session_state["plan_prompt"] = DEFAULT_PLAN_ANALYSIS_PROMPT
            st.rerun()
    
    with tabs[6]:
        st.subheader("Custom Rubric Upload")
        st.caption("Upload your institution's specific rubric to supplement default criteria.")
        
        rubric_file = st.file_uploader(
            "Upload rubric (PDF, Word, or Text)",
            type=["pdf", "docx", "txt"],
            key="rubric_upload"
        )
        
        if rubric_file:
            rubric_text = process_uploaded_file(rubric_file)
            if rubric_text:
                st.session_state["custom_rubric_text"] = rubric_text
                st.success(f"Loaded: {rubric_file.name}")
        
        st.divider()
        st.caption("Or paste rubric text:")
        st.session_state["custom_rubric_text"] = st.text_area(
            "Custom rubric text:",
            value=st.session_state.get("custom_rubric_text", ""),
            height=300,
            key="paste_rubric"
        )
        
        if st.session_state.get("custom_rubric_text"):
            if st.button("Clear Custom Rubric"):
                st.session_state["custom_rubric_text"] = ""
                st.rerun()
    
    with tabs[7]:
        st.subheader("Unit Registry Management")
        st.caption("View, edit, and add unit names. Changes are saved to session.")
        
        registry = load_unit_registry()
        
        # Display and edit registry
        reg_tab1, reg_tab2 = st.tabs(["Academic Units", "Administrative Units"])
        
        with reg_tab1:
            # Add new academic unit form
            with st.expander("➕ Add New Academic Unit", expanded=False):
                new_acad_col1, new_acad_col2 = st.columns(2)
                with new_acad_col1:
                    new_acad_name = st.text_input("Unit Name *", key="new_acad_name", placeholder="e.g., Computer Science")
                    new_acad_college = st.text_input("College/Department *", key="new_acad_college", placeholder="e.g., College of Engineering")
                with new_acad_col2:
                    new_acad_prev = st.text_input("Previous Names", key="new_acad_prev", placeholder="Semicolon-separated", help="e.g., CS Department; Comp Sci")
                    new_acad_active = st.selectbox("Active", ["Yes", "No"], key="new_acad_active")
                
                if st.button("Add Academic Unit", key="add_acad_btn"):
                    if new_acad_name and new_acad_college:
                        new_unit_id = generate_unit_id(new_acad_name, new_acad_college, "Academic")
                        new_unit = {
                            "unit_id": new_unit_id,
                            "canonical_name": new_acad_name,
                            "college_dept": new_acad_college,
                            "unit_type": "Academic",
                            "previous_names": new_acad_prev,
                            "active": new_acad_active
                        }
                        registry["academic"].append(new_unit)
                        st.session_state["unit_registry"] = registry
                        st.success(f"✓ Added: {new_acad_name} (ID: {new_unit_id})")
                        st.rerun()
                    else:
                        st.error("Unit Name and College/Department are required.")
            
            # Display existing units
            if registry["academic"]:
                df = pd.DataFrame(registry["academic"])
                # Ensure all columns are strings
                for col in df.columns:
                    df[col] = df[col].astype(str).replace('nan', '')
                
                st.caption(f"{len(df)} academic units")
                
                # Make editable
                edited_df = st.data_editor(
                    df,
                    use_container_width=True,
                    num_rows="fixed",
                    disabled=["unit_id", "unit_type"],
                    key="academic_editor"
                )
                
                if st.button("Save Changes to Academic Units", key="save_academic"):
                    # Regenerate unit_ids for any changed names
                    updated_units = []
                    for record in edited_df.to_dict('records'):
                        # Regenerate ID if name changed
                        new_id = generate_unit_id(record.get("canonical_name", ""), record.get("college_dept", ""), "Academic")
                        record["unit_id"] = new_id
                        record["unit_type"] = "Academic"
                        updated_units.append(record)
                    registry["academic"] = updated_units
                    st.session_state["unit_registry"] = registry
                    st.success("✓ Academic units saved!")
            else:
                st.info("No academic units loaded. Add one above or upload a CSV below.")
        
        with reg_tab2:
            # Add new administrative unit form
            with st.expander("➕ Add New Administrative Unit", expanded=False):
                new_admin_col1, new_admin_col2 = st.columns(2)
                with new_admin_col1:
                    new_admin_name = st.text_input("Unit Name *", key="new_admin_name", placeholder="e.g., Financial Aid")
                    new_admin_div = st.text_input("Division *", key="new_admin_div", placeholder="e.g., Student Affairs")
                with new_admin_col2:
                    new_admin_prev = st.text_input("Previous Names", key="new_admin_prev", placeholder="Semicolon-separated")
                    new_admin_active = st.selectbox("Active", ["Yes", "No"], key="new_admin_active")
                
                if st.button("Add Administrative Unit", key="add_admin_btn"):
                    if new_admin_name and new_admin_div:
                        new_unit_id = generate_unit_id(new_admin_name, new_admin_div, "Administrative")
                        new_unit = {
                            "unit_id": new_unit_id,
                            "canonical_name": new_admin_name,
                            "college_dept": new_admin_div,
                            "unit_type": "Administrative",
                            "previous_names": new_admin_prev,
                            "active": new_admin_active
                        }
                        registry["administrative"].append(new_unit)
                        st.session_state["unit_registry"] = registry
                        st.success(f"✓ Added: {new_admin_name} (ID: {new_unit_id})")
                        st.rerun()
                    else:
                        st.error("Unit Name and Division are required.")
            
            # Display existing units
            if registry["administrative"]:
                df = pd.DataFrame(registry["administrative"])
                # Ensure all columns are strings
                for col in df.columns:
                    df[col] = df[col].astype(str).replace('nan', '')
                
                st.caption(f"{len(df)} administrative units")
                
                # Make editable
                edited_df = st.data_editor(
                    df,
                    use_container_width=True,
                    num_rows="fixed",
                    disabled=["unit_id", "unit_type"],
                    key="admin_editor"
                )
                
                if st.button("Save Changes to Administrative Units", key="save_admin"):
                    # Regenerate unit_ids for any changed names
                    updated_units = []
                    for record in edited_df.to_dict('records'):
                        new_id = generate_unit_id(record.get("canonical_name", ""), record.get("college_dept", ""), "Administrative")
                        record["unit_id"] = new_id
                        record["unit_type"] = "Administrative"
                        updated_units.append(record)
                    registry["administrative"] = updated_units
                    st.session_state["unit_registry"] = registry
                    st.success("✓ Administrative units saved!")
            else:
                st.info("No administrative units loaded. Add one above or upload a CSV below.")
        
        st.divider()
        st.caption("**Bulk Upload:** CSV format should have columns: College/Division, Unit Name")
        
        col1, col2 = st.columns(2)
        with col1:
            academic_upload = st.file_uploader("Academic Units CSV", type=["csv"], key="academic_csv")
            if academic_upload:
                df = pd.read_csv(academic_upload)
                # Process into registry format
                new_academic = []
                for _, row in df.iterrows():
                    unit_id = generate_unit_id(
                        str(row.iloc[1]) if len(row) > 1 else "",
                        str(row.iloc[0]) if len(row) > 0 else "",
                        "Academic"
                    )
                    new_academic.append({
                        "unit_id": unit_id,
                        "canonical_name": str(row.iloc[1]) if len(row) > 1 else "",
                        "college_dept": str(row.iloc[0]) if len(row) > 0 else "",
                        "unit_type": "Academic",
                        "previous_names": "",
                        "active": "Yes"
                    })
                registry["academic"] = new_academic
                st.session_state["unit_registry"] = registry
                st.success(f"Loaded {len(new_academic)} academic units")
        
        with col2:
            admin_upload = st.file_uploader("Administrative Units CSV", type=["csv"], key="admin_csv")
            if admin_upload:
                df = pd.read_csv(admin_upload)
                new_admin = []
                for _, row in df.iterrows():
                    unit_id = generate_unit_id(
                        str(row.iloc[1]) if len(row) > 1 else "",
                        str(row.iloc[0]) if len(row) > 0 else "",
                        "Administrative"
                    )
                    new_admin.append({
                        "unit_id": unit_id,
                        "canonical_name": str(row.iloc[1]) if len(row) > 1 else "",
                        "college_dept": str(row.iloc[0]) if len(row) > 0 else "",
                        "unit_type": "Administrative",
                        "previous_names": "",
                        "active": "Yes"
                    })
                registry["administrative"] = new_admin
                st.session_state["unit_registry"] = registry
                st.success(f"Loaded {len(new_admin)} administrative units")

# ============================================================================
# MAIN APPLICATION
# ============================================================================

def main():
    """Main application entry point."""
    
    # Inject UTA branding
    inject_uta_branding()
    
    # Check authentication
    if not check_password():
        return
    
    # Load unit registry
    registry = load_unit_registry()
    
    # Initialize navigation state
    if "current_page" not in st.session_state:
        st.session_state["current_page"] = "analyze"
    
    # Sidebar
    with st.sidebar:
        # Logo area - clean, bold text
        st.markdown("""
        <div class="logo-area">
            <div class="logo-text">Assessment<br>Analyzer</div>
        </div>
        """, unsafe_allow_html=True)
        
        # Navigation section
        st.markdown('<p style="color: #8896a6; font-size: 0.7rem; font-weight: 600; text-transform: uppercase; letter-spacing: 0.5px; margin-bottom: 0.5rem;">Pages</p>', unsafe_allow_html=True)
        
        # Navigation buttons - use type="primary" for active page
        if st.session_state["current_page"] == "analyze":
            st.button("● Analyze Report", key="nav_analyze", use_container_width=True, type="primary", disabled=True)
        else:
            if st.button("Analyze Report", key="nav_analyze", use_container_width=True):
                st.session_state["current_page"] = "analyze"
                st.rerun()
        
        if st.session_state["current_page"] == "batch":
            st.button("● Batch Import", key="nav_batch", use_container_width=True, type="primary", disabled=True)
        else:
            if st.button("Batch Import", key="nav_batch", use_container_width=True):
                st.session_state["current_page"] = "batch"
                st.rerun()
        
        if st.session_state["admin_mode"]:
            if st.session_state["current_page"] == "config":
                st.button("● Configuration", key="nav_config", use_container_width=True, type="primary", disabled=True)
            else:
                if st.button("Configuration", key="nav_config", use_container_width=True):
                    st.session_state["current_page"] = "config"
                    st.rerun()
        
        st.divider()
        
        # Status section - only show full status for admin
        if st.session_state["admin_mode"]:
            st.markdown('<p style="color: #8896a6; font-size: 0.7rem; font-weight: 600; text-transform: uppercase; letter-spacing: 0.5px; margin-bottom: 0.5rem;">Status</p>', unsafe_allow_html=True)
        
        # Load credentials
        if st.session_state["admin_mode"]:
            with st.expander("API Configuration", expanded=False):
                api_key = st.text_input(
                    "Anthropic API Key",
                    value=os.environ.get("ANTHROPIC_API_KEY", ""),
                    type="password",
                    key="api_key_input"
                )
                
                st.markdown("---")
                st.markdown("**Excel Online**")
                
                ms_client_id = st.text_input("Client ID", value=os.environ.get("MS_CLIENT_ID", ""), type="password", key="ms_client_id")
                ms_client_secret = st.text_input("Client Secret", value=os.environ.get("MS_CLIENT_SECRET", ""), type="password", key="ms_client_secret")
                ms_tenant_id = st.text_input("Tenant ID", value=os.environ.get("MS_TENANT_ID", ""), key="ms_tenant_id")
                ms_drive_id = st.text_input("Drive ID", value=os.environ.get("MS_DRIVE_ID", ""), key="ms_drive_id")
                ms_item_id = st.text_input("Item ID", value=os.environ.get("MS_ITEM_ID", ""), key="ms_item_id")
        else:
            api_key = os.environ.get("ANTHROPIC_API_KEY", "")
            ms_client_id = os.environ.get("MS_CLIENT_ID", "")
            ms_client_secret = os.environ.get("MS_CLIENT_SECRET", "")
            ms_tenant_id = os.environ.get("MS_TENANT_ID", "")
            ms_drive_id = os.environ.get("MS_DRIVE_ID", "")
            ms_item_id = os.environ.get("MS_ITEM_ID", "")
        
        # Check Excel connection
        access_token = None
        excel_connected = False
        
        if ms_client_id and ms_client_secret and ms_tenant_id and ms_drive_id and ms_item_id:
            if ms_client_id != "pending" and ms_client_secret != "pending":
                access_token = get_graph_access_token(ms_client_id, ms_client_secret, ms_tenant_id)
                if access_token:
                    excel_connected = True
        
        # Only show connection status for admin
        if st.session_state["admin_mode"]:
            if excel_connected:
                st.markdown('<span class="status-pill status-connected">● Excel Connected</span>', unsafe_allow_html=True)
            elif ms_client_id and ms_client_id != "pending":
                st.markdown('<span class="status-pill status-pending">● Excel Error</span>', unsafe_allow_html=True)
            elif ms_client_id == "pending":
                st.markdown('<span class="status-pill status-pending">● Excel Pending</span>', unsafe_allow_html=True)
            else:
                st.markdown('<span class="status-pill status-disconnected">○ Excel Not Configured</span>', unsafe_allow_html=True)
        
        # User mode indicator
        st.markdown("<br>", unsafe_allow_html=True)
        if st.session_state["admin_mode"]:
            st.markdown('<span class="status-pill status-connected">Admin Mode</span>', unsafe_allow_html=True)
        else:
            st.markdown('<span class="status-pill status-disconnected">User Mode</span>', unsafe_allow_html=True)
        
        # Spacer
        st.markdown("<br><br><br>", unsafe_allow_html=True)
        
        # Logout at bottom
        st.divider()
        if st.button("Logout", key="logout_btn", use_container_width=True):
            st.session_state["authenticated"] = False
            st.session_state["admin_mode"] = False
            st.rerun()
    
    # Main content area based on current page
    if st.session_state["current_page"] == "analyze":
        render_analyze_page(api_key, access_token, ms_drive_id, ms_item_id, excel_connected, registry)
    elif st.session_state["current_page"] == "batch":
        render_batch_page(api_key, access_token, ms_drive_id, ms_item_id, excel_connected, registry)
    elif st.session_state["current_page"] == "config" and st.session_state["admin_mode"]:
        render_admin_panel()
    
    # Footer
    render_uta_footer()


def render_analyze_page(api_key, access_token, ms_drive_id, ms_item_id, excel_connected, registry):
    """Render the main analysis page with clean design."""
    
    render_uta_header("Analyze Report")
    
    # Two column layout
    col1, col2 = st.columns([1, 1], gap="large")
    
    with col1:
        # UPLOAD SECTION
        st.markdown('<p class="section-header">Upload</p>', unsafe_allow_html=True)
        
        report_type = st.selectbox(
            "Report Type",
            REPORT_TYPES,
            key="report_type"
        )
        
        uploaded_file = st.file_uploader(
            "Drop your report here or click to browse",
            type=["pdf", "docx", "txt"],
            key="single_upload"
        )
        
        if uploaded_file:
            st.success(f"✓ {uploaded_file.name}")
            
            with st.spinner("Extracting text..."):
                report_text = process_uploaded_file(uploaded_file)
            
            if report_text:
                with st.expander("Preview extracted text", expanded=False):
                    st.text(report_text[:3000] + "..." if len(report_text) > 3000 else report_text)
                
                st.markdown("<br>", unsafe_allow_html=True)
                
                if st.button("Analyze Report", type="primary", use_container_width=True):
                    if not api_key:
                        st.error("API key not configured.")
                    else:
                        # Extract metadata
                        with st.spinner("Extracting metadata..."):
                            metadata = extract_metadata_with_ai(report_text, report_type, api_key)
                        
                        if "error" in metadata:
                            st.error(f"Error: {metadata['error']}")
                        else:
                            st.session_state["extracted_metadata"] = metadata
                            st.session_state["filename"] = uploaded_file.name
                        
                        # Historical context
                        stagnation_info = None
                        previous_improvements = None
                        
                        if excel_connected and access_token and ms_drive_id and ms_item_id and metadata.get("unit_id"):
                            if report_type == "Results Report":
                                for outcome in metadata.get("outcomes", []):
                                    stag = check_stagnation(
                                        access_token, ms_drive_id, ms_item_id,
                                        metadata.get("unit_id", ""),
                                        outcome.get("outcome_id", ""),
                                        outcome.get("assessment_method", ""),
                                        metadata.get("academic_year", "")
                                    )
                                    if stag.get("stagnant"):
                                        stagnation_info = stag
                                        break
                            
                            elif report_type == "Improvement Report":
                                match = find_matching_unit(
                                    metadata.get("unit_name", ""),
                                    metadata.get("unit_type", "Academic"),
                                    registry
                                )
                                if match["match"]:
                                    previous_improvements = get_previous_improvements_excel(
                                        access_token, ms_drive_id, ms_item_id,
                                        match["match"].get("unit_id", "")
                                    )
                        
                        # Run analysis
                        with st.spinner("Analyzing..."):
                            results = analyze_report(
                                report_text, 
                                report_type, 
                                api_key,
                                stagnation_info,
                                previous_improvements
                            )
                        
                        if "error" in results:
                            st.error(f"Error: {results['error']}")
                        else:
                            st.session_state["results"] = results
    
    with col2:
        # RESULTS SECTION
        st.markdown('<p class="section-header">Analysis Results</p>', unsafe_allow_html=True)
        
        if st.session_state.get("results"):
            results = st.session_state["results"]
            
            st.caption(f"Cost: {results['cost']} · {results['tokens']['input']} in / {results['tokens']['output']} out tokens")
            
            # Formatted display
            st.markdown('<div class="results-card">', unsafe_allow_html=True)
            st.markdown(results["analysis"])
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Expandable copy section - plain text version
            with st.expander("📋 Copy to clipboard"):
                # Convert markdown to plain text
                plain_text = results["analysis"]
                plain_text = plain_text.replace("### ", "").replace("## ", "").replace("# ", "")
                plain_text = plain_text.replace("**", "").replace("*", "")
                plain_text = plain_text.replace("✓", "STRENGTHS:").replace("✎", "REVISIONS REQUESTED:")
                st.caption("Click the copy icon in the top-right corner of the box below:")
                st.code(plain_text, language=None)
        else:
            st.info("Upload a report and click 'Analyze Report' to see results.")
        
        # METADATA SECTION
        if st.session_state.get("extracted_metadata"):
            st.markdown("<br>", unsafe_allow_html=True)
            st.markdown('<p class="section-header">Extracted Metadata</p>', unsafe_allow_html=True)
            
            metadata = st.session_state["extracted_metadata"]
            report_type = st.session_state.get("report_type", "Results Report")
            
            edited_metadata = render_metadata_editor(metadata, report_type, registry)
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            col_a, col_b = st.columns(2)
            
            with col_a:
                if st.button("Save to Excel Online", type="primary", use_container_width=True):
                    if not excel_connected or not access_token:
                        st.error("Excel Online not connected.")
                    else:
                        rows = prepare_rows_for_sheet(edited_metadata, report_type)
                        if save_metadata_to_excel_online(access_token, ms_drive_id, ms_item_id, rows, report_type):
                            st.success("✓ Saved!")
                            st.session_state["extracted_metadata"] = None
                        else:
                            st.error("Save failed.")
            
            with col_b:
                rows = prepare_rows_for_sheet(edited_metadata, report_type)
                df = pd.DataFrame(rows)
                
                buffer = BytesIO()
                with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                    df.to_excel(writer, index=False)
                
                st.download_button(
                    "Download Excel",
                    data=buffer.getvalue(),
                    file_name=f"metadata_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )


def render_batch_page(api_key, access_token, ms_drive_id, ms_item_id, excel_connected, registry):
    """Render batch import page with clean design."""
    
    render_uta_header("Batch Import")
    
    st.info("Upload multiple historical reports to populate the metadata database. Analysis will not be performed.")
    
    st.markdown('<p class="section-header">Settings</p>', unsafe_allow_html=True)
    
    report_type = st.selectbox(
        "Report Type (all files must be same type)",
        REPORT_TYPES,
        key="batch_report_type"
    )
    
    st.markdown('<p class="section-header">Files</p>', unsafe_allow_html=True)
    
    uploaded_files = st.file_uploader(
        "Drop files here or click to browse",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        key="batch_files_upload"
    )
    
    if uploaded_files:
        st.success(f"✓ {len(uploaded_files)} files selected")
        
        if st.button("Extract Metadata from All Files", type="primary"):
            all_metadata = []
            progress = st.progress(0)
            
            for i, file in enumerate(uploaded_files):
                with st.spinner(f"Processing {file.name}..."):
                    text = process_uploaded_file(file)
                    if text:
                        metadata = extract_metadata_with_ai(text, report_type, api_key)
                        if "error" not in metadata:
                            metadata["_filename"] = file.name
                            all_metadata.append(metadata)
                        else:
                            st.warning(f"Error with {file.name}: {metadata['error']}")
                progress.progress((i + 1) / len(uploaded_files))
            
            st.session_state["batch_metadata"] = all_metadata
            st.success(f"✓ Extracted metadata from {len(all_metadata)} files")
        
        if st.session_state.get("batch_metadata"):
            st.markdown('<p class="section-header">Extracted Data</p>', unsafe_allow_html=True)
            
            for meta in st.session_state["batch_metadata"]:
                with st.expander(f"{meta.get('_filename', 'Unknown')} - {meta.get('unit_name', 'Unknown Unit')}"):
                    st.json(meta)
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("Save All to Excel Online", type="primary", use_container_width=True):
                    if not excel_connected or not access_token:
                        st.error("Excel Online not connected.")
                    else:
                        success_count = 0
                        for meta in st.session_state["batch_metadata"]:
                            rows = prepare_rows_for_sheet(meta, report_type)
                            if save_metadata_to_excel_online(access_token, ms_drive_id, ms_item_id, rows, report_type):
                                success_count += 1
                        st.success(f"✓ Saved {success_count} of {len(st.session_state['batch_metadata'])} records")

if __name__ == "__main__":
    main()